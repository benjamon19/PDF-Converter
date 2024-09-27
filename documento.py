import sys
import os
from datetime import datetime
from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, QSizePolicy, QSpacerItem, QPushButton, QMessageBox, QLineEdit, QListWidget, QFileDialog, QAction, QDialog, QTextEdit, QFormLayout, QComboBox, QDateEdit, QListWidgetItem, QInputDialog, QProgressBar)
from PyQt5.QtGui import QPixmap, QIcon, QTextCursor, QFontDatabase, QPainter, QPen, QFont, QImage
from PyQt5.QtCore import Qt, QDate, pyqtSignal
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from data.database import DatabaseManager

def resource_path(relative_path):
    """ Obtener la ruta absoluta al recurso, funciona para desarrollo y para PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def get_spanish_month_name(month_number):
    months = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    return months[month_number - 1]

class ModifyFieldsDialog(QDialog):
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db_manager = db_manager
        self.setWindowTitle("Modificar Campos")
        self.setMinimumWidth(450)
        self.fields = {}
        self.init_ui()
        self.load_styles()

    def init_ui(self):
        layout = QFormLayout(self)
        
        fields = [
            ("RUT", QLineEdit),
            ("NOMBRE", QLineEdit),
            ("ESTADO_CIVIL", QComboBox),
            ("FECHA_NACIMIENTO", QDateEdit),
            ("FONASA_ISAPRE", QComboBox),
            ("AFP", QComboBox),
            ("CONTACTO", QLineEdit),
            ("EMAIL", QLineEdit),
            ("PAIS", QLineEdit),
            ("NACIONALIDAD", QLineEdit),
            ("DOMICILIO", QLineEdit),
            ("CIUDAD", QLineEdit),
            ("SUELDO_BASE", QLineEdit),
            ("BONO_ASISTENCIA", QLineEdit),
            ("COLACION", QLineEdit),
            ("FECHA_CONTRATO", QDateEdit),
            ("DURACION_CONTRATO_INICIO", QDateEdit),
            ("DURACION_CONTRATO_FINAL", QDateEdit),
            ("TIPO_CONTRATO", QComboBox)
        ]

        estado_civil_options = ["Soltero/a", "Casado/a", "Viudo/a", "Divorciado/a"]
        fonasa_isapre_options = ["Fonasa", "Isapre"]
        afp_options = ["AFP Habitat", "AFP Capital", "AFP Cuprum", "AFP PlanVital", "AFP ProVida", "AFP Modelo"]
        tipo_contrato_options = ["Indefinido", "Plazo Fijo", "Por Obra o Faena", "Aprendizaje"]

        for field, widget_class in fields:
            label = QLabel(field.replace("_", " ").title())
            if widget_class == QComboBox:
                widget = widget_class(self)
                if field == "ESTADO_CIVIL":
                    widget.addItems(estado_civil_options)
                elif field == "FONASA_ISAPRE":
                    widget.addItems(fonasa_isapre_options)
                elif field == "AFP":
                    widget.addItems(afp_options)
                elif field == "TIPO_CONTRATO":
                    widget.addItems(tipo_contrato_options)
            elif widget_class == QDateEdit:
                widget = widget_class(self)
                widget.setCalendarPopup(True)
                widget.setDisplayFormat("yyyy-MM-dd")
            else:
                widget = widget_class(self)
                if field in ["SUELDO_BASE", "BONO_ASISTENCIA", "COLACION"]:
                    widget.setPlaceholderText("$")
            self.fields[field] = widget
            layout.addRow(label, widget)

        self.fields["RUT"].editingFinished.connect(self.populate_fields)

        self.save_button = QPushButton("Guardar", self)
        self.save_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.save_button.clicked.connect(self.save_fields)
        layout.addRow(self.save_button)

        self.setLayout(layout)

    def populate_fields(self):
        rut = self.fields["RUT"].text()
        try :
            if rut:
                worker_data = self.db_manager.get_contratista_by_rut(rut)
                if worker_data:
                    self.fields["NOMBRE"].setText(worker_data[2])
                    self.fields["ESTADO_CIVIL"].setCurrentText(worker_data[3])
                    self.fields["FECHA_NACIMIENTO"].setDate(QDate.fromString(worker_data[4], "yyyy-MM-dd"))
                    self.fields["FONASA_ISAPRE"].setCurrentText(worker_data[5])
                    self.fields["AFP"].setCurrentText(worker_data[6])
                    self.fields["CONTACTO"].setText(worker_data[7])
                    self.fields["EMAIL"].setText(worker_data[8])
                    self.fields["CIUDAD"].setText(worker_data[9])
                    self.fields["PAIS"].setText(worker_data[10])
                    self.fields["NACIONALIDAD"].setText(worker_data[11])
                    self.fields["DOMICILIO"].setText(worker_data[12])
                    self.save_button.setText("Modificar")
                else:
                    self.clear_fields()
        except : 
            pass

    def save_fields(self):
        fields_data = {
            "RUT": self.fields["RUT"].text(),
            "NOMBRE": self.fields["NOMBRE"].text(),
            "ESTADO_CIVIL": self.fields["ESTADO_CIVIL"].currentText(),
            "FECHA_NACIMIENTO": self.fields["FECHA_NACIMIENTO"].date().toString("yyyy-MM-dd"),
            "FONASA_ISAPRE": self.fields["FONASA_ISAPRE"].currentText(),
            "AFP": self.fields["AFP"].currentText(),
            "CONTACTO": self.fields["CONTACTO"].text(),
            "EMAIL": self.fields["EMAIL"].text(),
            "CIUDAD": self.fields["CIUDAD"].text(),
            "PAIS": self.fields["PAIS"].text(),
            "NACIONALIDAD": self.fields["NACIONALIDAD"].text(),
            "DOMICILIO": self.fields["DOMICILIO"].text()
        }

        try:
            self.db_manager.save_contratista(fields_data)
            self.fields_data = fields_data 
            self.accept()
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Éxito")
            msg_box.setText("Datos guardados exitosamente")
            msg_box.setIcon(QMessageBox.Information)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()
        except Exception as e:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Error")
            msg_box.setText(f"Error al guardar los datos: {str(e)}")
            msg_box.setIcon(QMessageBox.Critical)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()

    def load_styles(self):
        try:
            style_path = resource_path("styles/documento.qss")
            with open(style_path, "r") as file:
                self.setStyleSheet(file.read())
        except Exception as e:
            print(f"Error al cargar el archivo de estilo: {e}")

class FirmasDialog(QDialog):
    def __init__(self, db_manager, fields_data, template_path, parent=None):
        super().__init__(parent)
        self.db_manager = db_manager
        self.fields_data = fields_data
        self.template_path = template_path
        self.signature_path = None
        self.setWindowTitle("Firmas")
        self.setMinimumSize(600, 400)
        self.init_ui()
        self.load_signatures_from_db()
        self.load_styles()

    def init_ui(self):
        layout = QVBoxLayout(self)
        self.signature_list = QListWidget(self)
        layout.addWidget(self.signature_list)
        button_layout = QHBoxLayout()

        self.create_signature_button = QPushButton("Crear Firma", self)
        self.create_signature_button.clicked.connect(self.create_signature)
        button_layout.addWidget(self.create_signature_button)

        self.preview_signature_button = QPushButton("Vista Previa de Firma", self)
        self.preview_signature_button.clicked.connect(self.preview_signature)
        button_layout.addWidget(self.preview_signature_button)

        self.delete_signature_button = QPushButton("Eliminar Firma", self)
        self.delete_signature_button.clicked.connect(self.delete_signature)
        button_layout.addWidget(self.delete_signature_button)

        self.download_pdf_button = QPushButton("Descargar Contrato", self)
        self.download_pdf_button.clicked.connect(self.download_pdf)
        button_layout.addWidget(self.download_pdf_button)

        layout.addLayout(button_layout)

        self.setLayout(layout)

    def load_signatures_from_db(self):
        signatures = self.db_manager.load_signatures()
        self.signature_list.clear()
        for signature_id, name, path in signatures:
            signature_item = QListWidgetItem(name)
            signature_item.setData(Qt.UserRole, path)
            signature_item.setData(Qt.UserRole + 1, signature_id)
            self.signature_list.addItem(signature_item)

    def create_signature(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Crear Firma")

        layout = QVBoxLayout(dialog)
        label = QLabel("Ingrese el nombre para la firma:", dialog)
        layout.addWidget(label)

        line_edit = QLineEdit(dialog)
        layout.addWidget(line_edit)

        button_layout = QHBoxLayout()
        ok_button = QPushButton("OK", dialog)
        ok_button.setMinimumWidth(120)
        cancel_button = QPushButton("Cancel", dialog)
        cancel_button.setMinimumWidth(120)

        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        layout.addLayout(button_layout)

        ok_button.clicked.connect(dialog.accept)
        cancel_button.clicked.connect(dialog.reject)

        if dialog.exec_() == QDialog.Accepted:
            text = line_edit.text()
            if text:
                signature_dir = resource_path('firmas')
                if not os.path.exists(signature_dir):
                    os.makedirs(signature_dir)
                output_path = os.path.join(signature_dir, f"{text}.png")
                self.generate_signature(text, output_path)
                self.db_manager.save_signature(text, output_path)
                self.load_signatures_from_db()

    def generate_signature(self, name, output_path):
        width, height = 300, 100
        image = QImage(width, height, QImage.Format_ARGB32)
        image.fill(Qt.transparent)

        painter = QPainter(image)
        pen = QPen(Qt.black, 2)
        painter.setPen(pen)
        
        font_id = QFontDatabase.addApplicationFont(resource_path("fonts/GreatVibes-Regular.ttf"))
        if font_id == -1:
            print("Error al cargar la fuente.")
            return
        font_family = QFontDatabase.applicationFontFamilies(font_id)[0]
        font = QFont(font_family, 30, QFont.Normal)
        painter.setFont(font)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.setRenderHint(QPainter.TextAntialiasing)
        painter.drawText(image.rect(), Qt.AlignCenter, name)
        painter.end()
        image.save(output_path, "PNG")

    def preview_signature(self):
        selected_item = self.signature_list.currentItem()
        if selected_item:
            signature_path = selected_item.data(Qt.UserRole)
            if not os.path.exists(signature_path):
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("Advertencia")
                msg_box.setText(f"La firma no se encuentra en la ruta: {signature_path}")
                msg_box.setIcon(QMessageBox.Warning)
                ok_button = msg_box.addButton(QMessageBox.Ok)
                ok_button.setMinimumWidth(120)
                msg_box.exec_()
                return
            pixmap = QPixmap(signature_path)
            preview_dialog = QDialog(self)
            preview_dialog.setWindowTitle("Vista Previa de Firma")
            preview_dialog.setMinimumSize(300, 200)
            preview_dialog.setMaximumSize(300, 200)
            preview_layout = QVBoxLayout(preview_dialog)
            preview_label = QLabel()
            preview_label.setPixmap(pixmap.scaled(280, 180, Qt.KeepAspectRatio))
            preview_layout.addWidget(preview_label)
            preview_dialog.setLayout(preview_layout)
            preview_dialog.exec_()

    def delete_signature(self):
        selected_item = self.signature_list.currentItem()
        if not selected_item:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Advertencia")
            msg_box.setText("Seleccione una firma para eliminar.")
            msg_box.setIcon(QMessageBox.Warning)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()
            return

        signature_id = selected_item.data(Qt.UserRole + 1)
        signature_path = selected_item.data(Qt.UserRole)

        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Eliminar Firma")
        msg_box.setText(f"¿Está seguro de que desea eliminar la firma '{selected_item.text()}'?")
        msg_box.setIcon(QMessageBox.Question)
        yes_button = msg_box.addButton(QMessageBox.Yes)
        no_button = msg_box.addButton(QMessageBox.No)
        yes_button.setMinimumWidth(120)
        no_button.setMinimumWidth(120)
        msg_box.exec_()

        if msg_box.clickedButton() == yes_button:
            if self.db_manager.delete_signature(signature_id):
                if os.path.exists(signature_path):
                    os.remove(signature_path)
                self.load_signatures_from_db()
            else:
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("Error")
                msg_box.setText("Error al eliminar la firma de la base de datos.")
                msg_box.setIcon(QMessageBox.Critical)
                ok_button = msg_box.addButton(QMessageBox.Ok)
                ok_button.setMinimumWidth(120)
                msg_box.exec_()

    def modify_word_template(self, template_path, output_path, fields):
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            for key, value in fields.items():
                if key == "FIRMA":
                    continue
                if key in paragraph.text:
                    inline = paragraph.runs
                    for item in inline:
                        if key in item.text:
                            item.text = item.text.replace(key, value)
                            item.bold = True
        
        doc.save(output_path)

    def insert_signature_into_word(self, doc_path, signature_path):
        doc = Document(doc_path)
        signature_placeholder = "FIRMA"
        for paragraph in doc.paragraphs:
            if signature_placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(signature_placeholder, '')
                run = paragraph.add_run()
                run.add_picture(signature_path, width=Inches(2))
        doc.save(doc_path)

    def convert_to_pdf(self, word_path, pdf_path):
        convert(word_path, pdf_path)

    def download_pdf(self):
        options = QFileDialog.Options()
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar PDF", "", "Archivos PDF (*.pdf);;Todos los archivos (*)", options=options)
        if save_path:
            output_word_path = "Modified_Contract.docx"
            output_pdf_path = "Modified_Contract.pdf"

            current_date = datetime.now()
            formatted_date = f"{current_date.day} de {get_spanish_month_name(current_date.month)} del {current_date.year}"
            self.fields_data["FECHA"] = formatted_date

            self.modify_word_template(self.template_path, output_word_path, self.fields_data)
            if self.signature_list.currentItem():
                self.signature_path = self.signature_list.currentItem().data(Qt.UserRole)
                self.insert_signature_into_word(output_word_path, self.signature_path)
            self.convert_to_pdf(output_word_path, output_pdf_path)

            if os.path.exists(output_pdf_path):
                os.rename(output_pdf_path, save_path)
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("Éxito")
                msg_box.setText(f"PDF guardado con éxito en: {save_path}")
                msg_box.setIcon(QMessageBox.Information)
                ok_button = msg_box.addButton(QMessageBox.Ok)
                ok_button.setMinimumWidth(120)
                msg_box.exec_()

                if os.path.exists(output_word_path):
                    os.remove(output_word_path)
            else:
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("Error")
                msg_box.setText("No se encontró el archivo PDF generado.")
                msg_box.setIcon(QMessageBox.Warning)
                ok_button = msg_box.addButton(QMessageBox.Ok)
                ok_button.setMinimumWidth(120)
                msg_box.exec_()

            self.accept()

    def load_styles(self):
        try:
            style_path = resource_path("styles/documento.qss")
            with open(style_path, "r") as file:
                self.setStyleSheet(file.read())
        except Exception as e:
            print(f"Error al cargar el archivo de estilo: {e}")


class PreviewDialog(QDialog):
    def __init__(self, file_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Vista Previa")
        self.setMinimumSize(600, 400)
        self.init_ui(file_path)

    def init_ui(self, file_path):
        layout = QVBoxLayout(self)
        try:
            text_edit = QTextEdit(self)
            text_edit.setReadOnly(True)
            document = Document(file_path)
            for para in document.paragraphs:
                text_edit.append(para.text)
            text_edit.moveCursor(QTextCursor.Start)
            layout.addWidget(text_edit)
        except Exception as e:
            error_label = QLabel(f"Error al cargar el documento: {e}", self)
            layout.addWidget(error_label)
        self.setLayout(layout)

class GestionarDocumento(QWidget):
    def __init__(self, stacked_widget):
        super().__init__()
        self.stacked_widget = stacked_widget
        self.db_manager = DatabaseManager()
        self.selected_template_path = None
        self.init_ui()
        self.load_templates()
        self.load_styles()

    def load_styles(self):
        try:
            style_path = resource_path("styles/documento.qss")
            with open(style_path, "r") as file:
                self.setStyleSheet(file.read())
        except Exception as e:
            print(f"Error al cargar el archivo de estilo: {e}")

    def init_ui(self):
        self.setWindowTitle('Gestionar Documento')
        self.setGeometry(100, 100, 1200, 800)

        layout = QVBoxLayout(self)

        top_frame = QWidget(self)
        top_layout = QHBoxLayout(top_frame)
        top_layout.setSpacing(10)
        top_layout.setContentsMargins(10, 10, 10, 10)

        self.logo_label = QLabel(self)
        logo_pixmap = QPixmap(resource_path("images/logo.png"))
        self.logo_label.setPixmap(logo_pixmap.scaled(250, 250, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        self.logo_label.setAlignment(Qt.AlignRight)
        top_layout.addItem(QSpacerItem(20, 20, QSizePolicy.Expanding, QSizePolicy.Minimum))
        top_layout.addWidget(self.logo_label)
        layout.addWidget(top_frame)

        self.database_label = QLabel("Gestión de Documentos", self)
        self.database_label.setObjectName("databaseLabel")
        layout.addWidget(self.database_label, alignment=Qt.AlignLeft)

        search_layout = QHBoxLayout()
        self.search_input = QLineEdit(self)
        self.search_input.setObjectName("searchInput")
        self.search_input.setPlaceholderText("Buscar documento...")
        self.search_input.textChanged.connect(self.filter_table)

        search_icon_action = QAction(QIcon(resource_path("icons/search.svg")), "", self.search_input)
        self.search_input.addAction(search_icon_action, QLineEdit.LeadingPosition)
        search_layout.addWidget(self.search_input)
        search_layout.addItem(QSpacerItem(20, 20, QSizePolicy.Expanding, QSizePolicy.Minimum))

        layout.addLayout(search_layout)

        self.template_list = QListWidget(self)
        self.template_list.itemClicked.connect(self.display_selected_pdf)
        layout.addWidget(self.template_list)

        self.placeholder_label = QLabel("De momento está vacío, debe agregar plantillas mediante el botón Cargar Plantilla", self)
        self.placeholder_label.setAlignment(Qt.AlignCenter)
        self.placeholder_label.setObjectName("placeholderLabel")
        layout.addWidget(self.placeholder_label)

        bottom_frame = QWidget(self)
        bottom_layout = QHBoxLayout(bottom_frame)
        bottom_layout.setSpacing(10)
        bottom_layout.setContentsMargins(10, 10, 10, 10)

        self.back_button = QPushButton("Volver al menú", self)
        self.back_button.setObjectName("backButton")
        self.back_button.setIcon(QIcon(resource_path("icons/home_white.svg")))
        self.back_button.clicked.connect(self.go_to_menu)

        bottom_layout.addWidget(self.back_button, alignment=Qt.AlignLeft)
        bottom_layout.addItem(QSpacerItem(20, 20, QSizePolicy.Expanding, QSizePolicy.Minimum))

        self.load_template_button = QPushButton("Cargar Plantilla", self)
        self.load_template_button.setObjectName("loadTemplateButton")
        self.load_template_button.clicked.connect(self.load_template)

        self.preview_button = QPushButton("Vista Previa", self)
        self.preview_button.setObjectName("previewButton")
        self.preview_button.clicked.connect(self.preview_template)

        self.modify_template_button = QPushButton("Modificar Plantilla", self)
        self.modify_template_button.setObjectName("modifyTemplateButton")
        self.modify_template_button.clicked.connect(self.show_modify_fields_dialog)

        self.delete_template_button = QPushButton("Eliminar Plantilla", self)
        self.delete_template_button.setObjectName("deleteTemplateButton")
        self.delete_template_button.clicked.connect(self.delete_template)

        bottom_layout.addWidget(self.load_template_button)
        bottom_layout.addWidget(self.preview_button)
        bottom_layout.addWidget(self.modify_template_button)
        bottom_layout.addWidget(self.delete_template_button)

        layout.addWidget(bottom_frame)

        self.setLayout(layout)

    def go_to_menu(self):
        from LOGIC.navigation import Navigation
        navigator = Navigation(self.stacked_widget)
        navigator.show_menu()

    def load_template(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar Plantilla", "", "Documentos de Word (*.docx)", options=options)
        if file_path:
            file_name = os.path.basename(file_path)
            self.selected_template_path = file_path
            if self.db_manager.save_template(file_name, file_path):
                self.load_templates()

    def load_templates(self):
        self.template_list.clear()
        templates = self.db_manager.load_plantilla_data()
        for template in templates:
            item = QListWidgetItem(template[1])
            item.setData(Qt.UserRole, template[2])
            self.template_list.addItem(item)
        self.update_placeholder_visibility()

    def update_placeholder_visibility(self):
        if (self.template_list.count() == 0):
            self.placeholder_label.show()
        else:
            self.placeholder_label.hide()

    def filter_table(self):
        query = self.search_input.text().lower()
        for i in range(self.template_list.count()):
            item = self.template_list.item(i)
            item.setHidden(query not in item.text().lower())

    def preview_template(self):
        selected_items = self.template_list.selectedItems()
        if not selected_items:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Error")
            msg_box.setText("Seleccione una plantilla para previsualizar.")
            msg_box.setIcon(QMessageBox.Warning)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()
            return

        template_name = selected_items[0].text()
        templates = self.db_manager.load_plantilla_data()
        for template in templates:
            if template[1] == template_name:
                file_path = template[2]
                break

        if not file_path:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Error")
            msg_box.setText("No se pudo encontrar el archivo de la plantilla.")
            msg_box.setIcon(QMessageBox.Warning)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()
            return

        try:
            preview_dialog = PreviewDialog(file_path)
            preview_dialog.exec_()
        except Exception as e:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Error")
            msg_box.setText(f"Error al cargar el documento: {e}")
            msg_box.setIcon(QMessageBox.Critical)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()

    def delete_template(self):
        selected_items = self.template_list.selectedItems()
        if not selected_items:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Error")
            msg_box.setText("Seleccione una plantilla para eliminar.")
            msg_box.setIcon(QMessageBox.Warning)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()
            return

        template_name = selected_items[0].text()
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Eliminar Plantilla")
        msg_box.setText(f"¿Está seguro de que desea eliminar la plantilla '{template_name}'?")
        msg_box.setIcon(QMessageBox.Question)
        yes_button = msg_box.addButton(QMessageBox.Yes)
        no_button = msg_box.addButton(QMessageBox.No)
        yes_button.setMinimumWidth(120)
        no_button.setMinimumWidth(120)
        msg_box.exec_()

        if msg_box.clickedButton() == yes_button:
            if self.db_manager.delete_template(template_name):
                self.load_templates()
            else:
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("Error")
                msg_box.setText("Error al eliminar la plantilla.")
                msg_box.setIcon(QMessageBox.Critical)
                ok_button = msg_box.addButton(QMessageBox.Ok)
                ok_button.setMinimumWidth(120)
                msg_box.exec_()

    def display_selected_pdf(self, item):
        file_path = item.data(Qt.UserRole)
        self.selected_template_path = file_path

    def show_modify_fields_dialog(self):
        selected_items = self.template_list.selectedItems()
        if not selected_items:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Error")
            msg_box.setText("Seleccione una plantilla para modificar.")
            msg_box.setIcon(QMessageBox.Warning)
            ok_button = msg_box.addButton(QMessageBox.Ok)
            ok_button.setMinimumWidth(120)
            msg_box.exec_()
            return

        dialog = ModifyFieldsDialog(self.db_manager, self)
        if dialog.exec_() == QDialog.Accepted:
            firmas_dialog = FirmasDialog(self.db_manager, dialog.fields_data, self.selected_template_path)
            firmas_dialog.exec_()

if __name__ == '__main__':
    from PyQt5.QtWidgets import QApplication, QStackedWidget
    app = QApplication(sys.argv)
    stacked_widget = QStackedWidget()
    gestionar_documento_window = GestionarDocumento(stacked_widget)
    stacked_widget.addWidget(gestionar_documento_window)
    stacked_widget.setCurrentWidget(gestionar_documento_window)
    stacked_widget.show()
    sys.exit(app.exec_())